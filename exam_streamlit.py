import copy
import io
import os
import random
import re
import tempfile
import zipfile

import streamlit as st
from docx import Document
from lxml import etree


SET_LABELS = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _deep_copy_row(row):
    return copy.deepcopy(row._tr)


def _get_cell_text(row, col_idx: int) -> str:
    if col_idx >= len(row.cells):
        return ""
    return row.cells[col_idx].text.strip()


def _clean_preview_text(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^\s*\d+\s*[\.\)]\s*", "", text)
    return text


def _set_cell_text(tr_el, col_idx: int, new_text: str):
    tc_elements = tr_el.findall(f"{{{W_NS}}}tc")

    if col_idx >= len(tc_elements):
        return

    tc = tc_elements[col_idx]
    all_runs = tc.findall(f".//{{{W_NS}}}r")

    if not all_runs:
        para = tc.find(f".//{{{W_NS}}}p")
        if para is None:
            para = etree.SubElement(tc, f"{{{W_NS}}}p")
        r = etree.SubElement(para, f"{{{W_NS}}}r")
        t = etree.SubElement(r, f"{{{W_NS}}}t")
        t.text = new_text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        return

    first_run = all_runs[0]

    for t_el in first_run.findall(f"{{{W_NS}}}t"):
        first_run.remove(t_el)

    t_new = etree.SubElement(first_run, f"{{{W_NS}}}t")
    t_new.text = new_text
    t_new.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    for r_el in all_runs[1:]:
        parent = r_el.getparent()
        if parent is not None:
            parent.remove(r_el)


def _strip_leading_number_from_cell(tr_el, col_idx: int):
    """
    Removes old question numbers from the question-body cell.
    Example:
    '14. What is ...' becomes 'What is ...'

    This prevents output like:
    13. 14. What is ...
    """
    tc_elements = tr_el.findall(f"{{{W_NS}}}tc")

    if col_idx >= len(tc_elements):
        return

    tc = tc_elements[col_idx]
    text_nodes = tc.findall(f".//{{{W_NS}}}t")

    if not text_nodes:
        return

    combined = "".join(t.text or "" for t in text_nodes)
    match = re.match(r"^\s*\d+\s*[\.\)]\s*", combined)

    if not match:
        return

    chars_to_remove = len(match.group(0))

    for t in text_nodes:
        current = t.text or ""

        if chars_to_remove <= 0:
            break

        if len(current) <= chars_to_remove:
            chars_to_remove -= len(current)
            t.text = ""
        else:
            t.text = current[chars_to_remove:]
            chars_to_remove = 0

        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def collect_questions(src_bytes: bytes):
    """
    Do NOT use @st.cache_data here.
    XML objects cannot be safely cached/pickled by Streamlit.
    """
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(src_bytes)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)
        questions = []

        for table in doc.tables:
            for row in table.rows:
                pts = _get_cell_text(row, 2)
                preview = _get_cell_text(row, 1)[:160].replace("\n", " ")
                preview = _clean_preview_text(preview)

                tr_copy = _deep_copy_row(row)

                questions.append(
                    {
                        "tr": tr_copy,
                        "pts": pts,
                        "preview": preview,
                    }
                )

    finally:
        os.unlink(tmp_path)

    return questions


def _patch_set_label(docx_bytes: bytes, new_label: str) -> bytes:
    pattern = re.compile(rb"Set [A-Z]")
    replacement = f"Set {new_label}".encode()

    buf_in = io.BytesIO(docx_bytes)
    buf_out = io.BytesIO()

    with zipfile.ZipFile(buf_in, "r") as zin, zipfile.ZipFile(
        buf_out, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        for name in zin.namelist():
            data = zin.read(name)

            if "footer" in name or "header" in name:
                data = pattern.sub(replacement, data)

            zout.writestr(name, data)

    return buf_out.getvalue()


def generate_set_bytes(src_bytes: bytes, questions: list, seed: int, label: str) -> bytes:
    shuffled = questions[:]
    random.Random(seed).shuffle(shuffled)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(src_bytes)
        src_tmp = tmp.name

    try:
        doc = Document(src_tmp)
        existing_rows = [row._tr for table in doc.tables for row in table.rows]

        if len(shuffled) != len(existing_rows):
            raise ValueError(
                f"Row count mismatch: document has {len(existing_rows)} rows, "
                f"but {len(shuffled)} were extracted."
            )

        for q_idx, (orig_tr, q) in enumerate(zip(existing_rows, shuffled)):
            parent = orig_tr.getparent()
            pos = list(parent).index(orig_tr)

            new_tr = copy.deepcopy(q["tr"])

            # Correct serial number in column 0
            _set_cell_text(new_tr, 0, str(q_idx + 1))

            # Remove old serial number from question body column
            _strip_leading_number_from_cell(new_tr, 1)

            # Keep original point value
            _set_cell_text(new_tr, 2, q["pts"])

            parent.remove(orig_tr)
            parent.insert(pos, new_tr)

        out_buf = io.BytesIO()
        doc.save(out_buf)
        raw_bytes = out_buf.getvalue()

    finally:
        os.unlink(src_tmp)

    return _patch_set_label(raw_bytes, label)


st.set_page_config(
    page_title="Exam Set Generator",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded",
)


st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@600;700&family=Source+Sans+3:wght@300;400;600&display=swap');

:root {
    --navy: #0d1b2a;
    --ink: #162032;
    --slate: #1c2f48;
    --gold: #d4a23a;
    --gold2: #f0c96a;
    --cream: #f2ece0;
    --muted: #7a93b0;
    --border: rgba(212,162,58,.25);
    --radius: 12px;
}

[data-testid="stAppViewContainer"],
[data-testid="stAppViewContainer"] > .main {
    background: var(--navy) !important;
}

html, body, p, li, span, label, div {
    font-family: 'Source Sans 3', sans-serif !important;
    color: var(--cream);
}

h1, h2, h3, h4 {
    font-family: 'Playfair Display', serif !important;
    color: var(--gold2) !important;
}

[data-testid="stSidebar"] {
    background: var(--ink) !important;
    border-right: 1px solid var(--border) !important;
}

[data-testid="stSidebar"] * {
    color: var(--cream) !important;
}

[data-testid="stFileUploader"] {
    border: 2px dashed var(--gold) !important;
    border-radius: var(--radius) !important;
    background: rgba(28,47,72,.5) !important;
}

.stDownloadButton > button,
.stButton > button {
    background: linear-gradient(135deg, var(--gold), #a8782a) !important;
    color: var(--navy) !important;
    font-weight: 700 !important;
    border: none !important;
    border-radius: 8px !important;
}

[data-testid="stMetric"] {
    background: var(--slate);
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: .8rem 1rem;
}

[data-testid="stMetricValue"] {
    color: var(--gold2) !important;
}

.q-card {
    display: flex;
    align-items: flex-start;
    gap: .75rem;
    background: rgba(255,255,255,.03);
    border-left: 3px solid var(--gold);
    border-radius: 0 6px 6px 0;
    padding: .6rem .9rem;
    margin-bottom: .35rem;
}

.q-num {
    font-family:'Playfair Display',serif;
    color:var(--gold);
    min-width:1.8rem;
}

.q-body {
    color:var(--cream);
    font-size:.9rem;
    line-height:1.5;
    flex:1;
}

.q-pts {
    color:var(--muted);
    font-size:.8rem;
    white-space:nowrap;
}
</style>
""",
    unsafe_allow_html=True,
)


with st.sidebar:
    st.markdown("## ⚙️ Configuration")
    st.markdown("---")

    num_sets = st.slider(
        "Number of exam sets",
        min_value=2,
        max_value=10,
        value=5,
        step=1,
    )

    st.info(f"Generating {num_sets} sets: {' · '.join(SET_LABELS[:num_sets])}")

    st.markdown("---")
    st.markdown(
        """
        **Expected Word file format**

        Your `.docx` should contain questions inside a table:

        - Column 0 = question number  
        - Column 1 = question body and choices  
        - Column 2 = point value  

        Important: If column 1 also starts with old numbers like `14.`, this app removes them automatically.
        """
    )


st.markdown(
    """
<div style='background:linear-gradient(135deg,#1c2f48 0%,#0d1b2a 100%);
            border:1px solid rgba(212,162,58,.3);
            border-left:5px solid #d4a23a;
            border-radius:14px;
            padding:1.8rem 2.2rem;
            margin-bottom:1.6rem;'>
    <h1 style='margin:0 0 .3rem;font-size:2.1rem;'>🎓 Exam Set Generator</h1>
    <p style='color:#7a93b0;margin:0;font-size:.95rem;'>
        Upload a Word exam document and generate multiple uniquely shuffled,
        print-ready sets with corrected serial numbers.
    </p>
</div>
""",
    unsafe_allow_html=True,
)


st.markdown("### 📂 Upload Exam Document")

uploaded = st.file_uploader(
    "Drop your .docx exam file here",
    type=["docx"],
    label_visibility="collapsed",
)


if uploaded is not None:
    src_bytes = uploaded.read()

    with st.spinner("Parsing questions from document..."):
        try:
            questions = collect_questions(src_bytes)
        except Exception as exc:
            st.error(f"Could not parse file: {exc}")
            st.stop()

    if not questions:
        st.warning("No table rows found. Make sure questions are inside a Word table.")
        st.stop()

    total_pts = sum(int(q["pts"]) for q in questions if q["pts"].isdigit())

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Questions", len(questions))
    c2.metric("Total Points", total_pts)
    c3.metric("Sets", num_sets)
    c4.metric("Method", "Seed Shuffle")

    st.success(f"{uploaded.name} loaded successfully. {len(questions)} questions detected.")

    st.markdown("---")
    st.markdown("### 📋 Generated Question Sets")

    tab_labels = [f"Set {SET_LABELS[i]}" for i in range(num_sets)]
    tabs = st.tabs(tab_labels)

    for i, tab in enumerate(tabs):
        label = SET_LABELS[i]
        seed = 1000 + i * 97

        shuffled = questions[:]
        random.Random(seed).shuffle(shuffled)

        with tab:
            q_cards = ""

            for idx, q in enumerate(shuffled, 1):
                q_cards += f"""
                <div class='q-card'>
                    <span class='q-num'>{idx}.</span>
                    <span class='q-body'>{q["preview"]}...</span>
                    <span class='q-pts'>{q["pts"]} pts</span>
                </div>
                """

            st.markdown(q_cards, unsafe_allow_html=True)

            with st.spinner(f"Building Set {label}..."):
                try:
                    docx_bytes = generate_set_bytes(src_bytes, questions, seed, label)
                except Exception as exc:
                    st.error(f"Error generating Set {label}: {exc}")
                    continue

            base_name = os.path.splitext(uploaded.name)[0]

            st.download_button(
                label=f"⬇️ Download Set {label} (.docx)",
                data=docx_bytes,
                file_name=f"{base_name}_Set_{label}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_{label}",
                use_container_width=True,
            )

else:
    st.info("Upload a `.docx` file to start generating exam sets.")


st.markdown("---")
st.caption("Exam Set Generator · Streamlit + python-docx")